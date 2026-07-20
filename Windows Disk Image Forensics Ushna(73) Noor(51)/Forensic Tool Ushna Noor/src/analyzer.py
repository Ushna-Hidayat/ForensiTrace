# ============================================================
#  analyzer.py  –  ForensiTrace Core Analysis Engine
#  Course  : CSDF-30117  Introduction to Digital Forensics
#  Author  : [Your Name]  |  Student ID: [Your ID]
#  Version : 1.0  |  Spring 2026
# ============================================================

import os
import csv
import hashlib
import datetime
import platform
import mimetypes
import struct
import json
from pathlib import Path

# ─────────────────────────────────────────────
#  1.  CSV / Timeline Analysis
# ─────────────────────────────────────────────

class TimelineAnalyzer:
    """Parses Autopsy-generated CSV and extracts forensic insights."""

    SUSPICIOUS_EXTENSIONS = {
        '.exe', '.bat', '.cmd', '.ps1', '.vbs', '.js', '.scr',
        '.dll', '.msi', '.com', '.pif', '.lnk', '.jar', '.hta'
    }
    SENSITIVE_KEYWORDS = [
        'password', 'passwd', 'secret', 'credential', 'login',
        'private', 'key', 'token', 'confidential', 'classified'
    ]

    def __init__(self, csv_path: str):
        self.csv_path   = csv_path
        self.records    = []
        self.suspicious = []
        self.anomalies  = []
        self.stats      = {}

    # ── Load ──────────────────────────────────
    def load(self) -> bool:
        """Read CSV into memory. Returns True on success."""
        try:
            with open(self.csv_path, newline='', encoding='utf-8', errors='replace') as f:
                reader = csv.DictReader(f)
                self.records = [row for row in reader]
            return True
        except Exception as e:
            print(f"[ERROR] Failed to load CSV: {e}")
            return False

    # ── Analyse ───────────────────────────────
    def analyse(self) -> dict:
        """Run all forensic checks. Returns summary dict."""
        if not self.records:
            return {}

        self.suspicious = []
        self.anomalies  = []

        for idx, row in enumerate(self.records):
            row_str = ' '.join(str(v) for v in row.values()).lower()

            # Suspicious extension check
            for ext in self.SUSPICIOUS_EXTENSIONS:
                # look in any field that might hold a filename
                for field in ('Name', 'File Name', 'name', 'filename', 'Path'):
                    val = row.get(field, '')
                    if val and val.lower().endswith(ext):
                        entry = dict(row)
                        entry['_reason'] = f'Suspicious extension: {ext}'
                        entry['_row']    = idx + 2
                        self.suspicious.append(entry)
                        break

            # Sensitive-keyword check
            for kw in self.SENSITIVE_KEYWORDS:
                if kw in row_str:
                    entry = dict(row)
                    entry['_reason'] = f'Sensitive keyword: "{kw}"'
                    entry['_row']    = idx + 2
                    if entry not in self.anomalies:
                        self.anomalies.append(entry)
                    break

        # Build stats
        cols = list(self.records[0].keys()) if self.records else []
        self.stats = {
            'total_records'   : len(self.records),
            'suspicious_files': len(self.suspicious),
            'anomalies'       : len(self.anomalies),
            'columns'         : cols,
            'csv_path'        : self.csv_path,
        }
        return self.stats

    # ── Timeline ──────────────────────────────
    def build_timeline(self, date_field: str | None = None) -> list[dict]:
        """Return records sorted by the best available date column."""
        date_candidates = ['Date', 'date', 'Timestamp', 'timestamp',
                           'Modified', 'Created', 'Accessed', 'Time']
        if date_field:
            date_candidates.insert(0, date_field)

        chosen = None
        for c in date_candidates:
            if self.records and c in self.records[0]:
                chosen = c
                break

        if not chosen:
            return self.records  # return as-is

        def _parse(val):
            for fmt in ('%Y-%m-%d %H:%M:%S', '%Y-%m-%dT%H:%M:%S',
                        '%m/%d/%Y %H:%M', '%Y-%m-%d', '%d/%m/%Y'):
                try:
                    return datetime.datetime.strptime(str(val).strip(), fmt)
                except ValueError:
                    pass
            return datetime.datetime.min

        return sorted(self.records, key=lambda r: _parse(r.get(chosen, '')))


# ─────────────────────────────────────────────
#  2.  File / Image Metadata Extractor
# ─────────────────────────────────────────────

class FileMetadataExtractor:
    """Extracts deep forensic metadata from any file on disk."""

    # JPEG / PNG magic bytes
    MAGIC = {
        b'\xff\xd8\xff'    : 'JPEG Image',
        b'\x89PNG\r\n\x1a\n': 'PNG Image',
        b'GIF87a'          : 'GIF Image',
        b'GIF89a'          : 'GIF Image',
        b'BM'              : 'BMP Image',
        b'%PDF'            : 'PDF Document',
        b'PK\x03\x04'      : 'ZIP / Office Document',
        b'\x1f\x8b'        : 'GZIP Archive',
        b'MZ'              : 'Windows PE Executable',
        b'\x7fELF'         : 'ELF Executable (Linux)',
    }

    def extract(self, file_path: str) -> dict:
        """Return a dict of forensic metadata for the given file."""
        p    = Path(file_path)
        info = {
            'file_name'   : p.name,
            'file_path'   : str(p.resolve()),
            'extension'   : p.suffix.lower(),
            'size_bytes'  : 0,
            'size_human'  : 'N/A',
            'mime_type'   : mimetypes.guess_type(str(p))[0] or 'unknown',
            'magic_type'  : 'Unknown',
            'md5'         : 'N/A',
            'sha256'      : 'N/A',
            'created'     : 'N/A',
            'modified'    : 'N/A',
            'accessed'    : 'N/A',
            'is_hidden'   : p.name.startswith('.'),
            'is_readonly' : False,
            'exif'        : {},
            'warnings'    : [],
        }

        if not p.exists():
            info['warnings'].append('File does not exist.')
            return info

        # ── Basic stat ────────────────────────
        st = p.stat()
        info['size_bytes'] = st.st_size
        info['size_human'] = self._human(st.st_size)
        info['modified']   = datetime.datetime.fromtimestamp(st.st_mtime).strftime('%Y-%m-%d %H:%M:%S')
        info['accessed']   = datetime.datetime.fromtimestamp(st.st_atime).strftime('%Y-%m-%d %H:%M:%S')
        if platform.system() == 'Windows':
            info['created'] = datetime.datetime.fromtimestamp(st.st_ctime).strftime('%Y-%m-%d %H:%M:%S')
        else:
            info['created'] = info['modified']  # Linux has no birth time in all FSes

        # ── Read-only flag ────────────────────
        import stat as stat_mod
        info['is_readonly'] = not os.access(str(p), os.W_OK)

        # ── Magic bytes ───────────────────────
        try:
            with open(str(p), 'rb') as f:
                header = f.read(16)
            for magic, label in self.MAGIC.items():
                if header.startswith(magic):
                    info['magic_type'] = label
                    break
        except Exception:
            pass

        # ── Hashes ───────────────────────────
        try:
            md5    = hashlib.md5()
            sha256 = hashlib.sha256()
            with open(str(p), 'rb') as f:
                for chunk in iter(lambda: f.read(65536), b''):
                    md5.update(chunk)
                    sha256.update(chunk)
            info['md5']    = md5.hexdigest()
            info['sha256'] = sha256.hexdigest()
        except Exception as e:
            info['warnings'].append(f'Hash error: {e}')

        # ── EXIF (images) ─────────────────────
        if info['extension'] in ('.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.gif', '.webp'):
            info['exif'] = self._extract_exif(str(p))

        # ── Suspicion flags ───────────────────
        if info['extension'] in TimelineAnalyzer.SUSPICIOUS_EXTENSIONS:
            info['warnings'].append('⚠ Suspicious file type detected.')
        if info['size_bytes'] == 0:
            info['warnings'].append('⚠ Zero-byte file — possible deletion artefact.')

        return info

    # ── EXIF via Pillow ───────────────────────
    def _extract_exif(self, path: str) -> dict:
        try:
            from PIL import Image
            from PIL.ExifTags import TAGS, GPSTAGS
            img  = Image.open(path)
            data = {}

            # Basic image info
            data['image_format'] = img.format or 'N/A'
            data['image_mode']   = img.mode
            data['image_size']   = f'{img.width} × {img.height} px'

            # Raw EXIF
            raw = img._getexif() if hasattr(img, '_getexif') else None
            if raw:
                for tag_id, value in raw.items():
                    tag = TAGS.get(tag_id, str(tag_id))
                    if tag == 'GPSInfo' and isinstance(value, dict):
                        gps = {}
                        for k, v in value.items():
                            gps[GPSTAGS.get(k, k)] = str(v)
                        data['GPS'] = gps
                    elif isinstance(value, bytes):
                        try:
                            data[tag] = value.decode('utf-8', errors='replace')
                        except Exception:
                            data[tag] = repr(value)
                    else:
                        data[tag] = str(value)
            return data
        except Exception as e:
            return {'error': str(e)}

    # ── Helper ────────────────────────────────
    @staticmethod
    def _human(n: int) -> str:
        for unit in ('B', 'KB', 'MB', 'GB', 'TB'):
            if n < 1024:
                return f'{n:.1f} {unit}'
            n /= 1024
        return f'{n:.1f} PB'


# ─────────────────────────────────────────────
#  3.  Hash Verifier
# ─────────────────────────────────────────────

class HashVerifier:
    """Compute and compare file hashes for integrity verification."""

    @staticmethod
    def compute(file_path: str) -> dict:
        algos = {
            'MD5'   : hashlib.md5(),
            'SHA1'  : hashlib.sha1(),
            'SHA256': hashlib.sha256(),
        }
        try:
            with open(file_path, 'rb') as f:
                for chunk in iter(lambda: f.read(65536), b''):
                    for h in algos.values():
                        h.update(chunk)
            return {name: h.hexdigest() for name, h in algos.items()}
        except Exception as e:
            return {'error': str(e)}

    @staticmethod
    def verify(file_path: str, expected_hash: str, algorithm: str = 'SHA256') -> dict:
        computed = HashVerifier.compute(file_path)
        if 'error' in computed:
            return {'match': False, 'reason': computed['error']}
        actual = computed.get(algorithm.upper(), '')
        match  = actual.lower() == expected_hash.lower().strip()
        return {
            'match'    : match,
            'algorithm': algorithm.upper(),
            'expected' : expected_hash,
            'actual'   : actual,
        }


# ─────────────────────────────────────────────
#  4.  Report Builder (PDF)
# ─────────────────────────────────────────────

class ForensicReportBuilder:
    """Generates a professional PDF forensic investigation report."""

    def __init__(self, output_path: str):
        self.output_path  = output_path
        self.sections     = []   # list of (heading, content_list, table_data)
        self.title        = 'Forensic Investigation Report'
        self.investigator = ''
        self.case_id      = ''
        self.generated_at = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        self.timeline_records = []   # full list of dicts for table rendering

    def set_meta(self, title: str, investigator: str, case_id: str):
        self.title        = title
        self.investigator = investigator
        self.case_id      = case_id

    def set_timeline_records(self, records: list[dict]):
        """Pass ALL timeline records for full table rendering in report."""
        self.timeline_records = records

    def add_section(self, heading: str, content: list[str]):
        """Add a text-only section."""
        self.sections.append((heading, content, None))

    def add_table_section(self, heading: str, content: list[str], table_data: list[list]):
        """Add a section that includes a data table."""
        self.sections.append((heading, content, table_data))

    def build(self) -> bool:
        try:
            from reportlab.lib.pagesizes import A4, landscape
            from reportlab.lib import colors
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.platypus import (SimpleDocTemplate, Paragraph,
                                            Spacer, Table, TableStyle,
                                            HRFlowable, PageBreak, KeepTogether)

            # Use landscape for timeline table (wider)
            page_size = landscape(A4) if self.timeline_records else A4

            doc = SimpleDocTemplate(self.output_path, pagesize=page_size,
                                    leftMargin=1.5*cm, rightMargin=1.5*cm,
                                    topMargin=2*cm, bottomMargin=2*cm)
            styles = getSampleStyleSheet()

            # ── Colours ──────────────────────────────────────
            DARK   = colors.HexColor('#0d1b2a')
            ACCENT = colors.HexColor('#1a73e8')
            ACCENT2= colors.HexColor('#0a4fa3')
            LIGHT  = colors.HexColor('#f4f6fa')
            GREEN  = colors.HexColor('#1a7a4a')
            RED    = colors.HexColor('#c0392b')

            # ── Styles ───────────────────────────────────────
            univ_style  = ParagraphStyle('FUniv', parent=styles['Normal'],
                                          fontSize=11, textColor=DARK,
                                          fontName='Helvetica-Bold',
                                          alignment=1, spaceAfter=2)
            dept_style  = ParagraphStyle('FDept', parent=styles['Normal'],
                                          fontSize=9, textColor=colors.grey,
                                          alignment=1, spaceAfter=2)
            title_style = ParagraphStyle('FTitle', parent=styles['Title'],
                                          fontSize=20, textColor=ACCENT,
                                          spaceAfter=4, fontName='Helvetica-Bold',
                                          alignment=1)
            label_style = ParagraphStyle('FLabel', parent=styles['Normal'],
                                          fontSize=9, textColor=DARK,
                                          fontName='Helvetica-Bold')
            h1_style    = ParagraphStyle('FH1', parent=styles['Heading1'],
                                          fontSize=13, textColor=ACCENT,
                                          fontName='Helvetica-Bold',
                                          spaceBefore=14, spaceAfter=4)
            body_style  = ParagraphStyle('FBody', parent=styles['Normal'],
                                          fontSize=9, leading=14, spaceAfter=3)
            warn_style  = ParagraphStyle('FWarn', parent=styles['Normal'],
                                          fontSize=9, textColor=RED, leading=14)
            cell_style  = ParagraphStyle('FCell', parent=styles['Normal'],
                                          fontSize=7, leading=9)

            story = []

            # ════════════════════════════════════════════════
            #  COVER PAGE
            # ════════════════════════════════════════════════
            story.append(Spacer(1, 0.8*cm))

            # University header
            story.append(Paragraph('THE ISLAMIA UNIVERSITY OF BAHAWALPUR', univ_style))
            story.append(Paragraph('Faculty of Engineering & Technology', dept_style))
            story.append(Paragraph('Department of Information and Communication Engineering', dept_style))
            story.append(Spacer(1, 0.3*cm))
            story.append(HRFlowable(width='100%', thickness=2, color=ACCENT))
            story.append(Spacer(1, 0.5*cm))

            # Report title
            story.append(Paragraph('FORENSIC INVESTIGATION REPORT', title_style))
            story.append(Paragraph('ForensiTrace — Digital Evidence Analysis &amp; Timeline Reconstruction',
                                   ParagraphStyle('sub', parent=styles['Normal'],
                                                  fontSize=10, textColor=colors.grey,
                                                  alignment=1)))
            story.append(Spacer(1, 0.6*cm))
            story.append(HRFlowable(width='100%', thickness=1, color=colors.lightgrey))
            story.append(Spacer(1, 0.5*cm))

            # Submission info table
            sub_data = [
                ['Submitted To',  'Engr. Rafia Durrani'],
                ['Course',        'CSDF-30117  Introduction to Digital Forensics'],
                ['Submitted By',  'Ushna Hidayat  |  Noor Fatima'],
                ['Case ID',       self.case_id      or 'N/A'],
                ['Report Title',  self.title],
                ['Generated At',  self.generated_at],
                ['Semester',      'Spring 2026'],
                ['Classification','CONFIDENTIAL — FOR ACADEMIC USE ONLY'],
            ]
            sub_table = Table(sub_data, colWidths=[4*cm, 14*cm])
            sub_table.setStyle(TableStyle([
                ('BACKGROUND',    (0, 0), (0, -1), LIGHT),
                ('TEXTCOLOR',     (0, 0), (0, -1), DARK),
                ('FONTNAME',      (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE',      (0, 0), (-1, -1), 9),
                ('ROWBACKGROUNDS',(0, 0), (-1, -1), [colors.white, LIGHT]),
                ('GRID',          (0, 0), (-1, -1), 0.4, colors.lightgrey),
                ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
                ('TOPPADDING',    (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                # Highlight submitted-to and submitted-by rows
                ('BACKGROUND',    (1, 0), (1, 0), colors.HexColor('#dbeafe')),
                ('BACKGROUND',    (1, 2), (1, 2), colors.HexColor('#dcfce7')),
                ('FONTNAME',      (1, 2), (1, 2), 'Helvetica-Bold'),
                ('TEXTCOLOR',     (1, 2), (1, 2), GREEN),
                ('FONTNAME',      (1, 0), (1, 0), 'Helvetica-Bold'),
                ('TEXTCOLOR',     (1, 0), (1, 0), ACCENT2),
            ]))
            story.append(sub_table)
            story.append(PageBreak())

            # ════════════════════════════════════════════════
            #  FULL TIMELINE TABLE  (all records)
            # ════════════════════════════════════════════════
            if self.timeline_records:
                story.append(Paragraph('Timeline Analysis — Complete Record Log', h1_style))
                story.append(HRFlowable(width='100%', thickness=0.5, color=colors.lightgrey))
                story.append(Spacer(1, 0.2*cm))
                story.append(Paragraph(
                    f'Total records exported from Autopsy CSV: {len(self.timeline_records)}',
                    body_style))
                story.append(Spacer(1, 0.3*cm))

                # Build table data
                cols = [c for c in self.timeline_records[0].keys() if not c.startswith('_')]
                # Limit columns to keep readable (take first 8 most useful)
                cols = cols[:8]

                hdr  = [Paragraph(f'<b>{c}</b>', cell_style) for c in cols]
                rows = [hdr]
                for rec in self.timeline_records:
                    row = []
                    for c in cols:
                        val = str(rec.get(c, ''))[:80]   # truncate very long values
                        row.append(Paragraph(val, cell_style))
                    rows.append(row)

                # Dynamic column width
                avail_w = 25*cm   # landscape A4 usable
                col_w   = avail_w / len(cols)
                tbl = Table(rows, colWidths=[col_w]*len(cols), repeatRows=1)
                tbl.setStyle(TableStyle([
                    ('BACKGROUND',    (0, 0), (-1, 0), ACCENT),
                    ('TEXTCOLOR',     (0, 0), (-1, 0), colors.white),
                    ('FONTNAME',      (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE',      (0, 0), (-1, -1), 7),
                    ('ROWBACKGROUNDS',(0, 1), (-1, -1),
                     [colors.white, colors.HexColor('#f0f4ff')]),
                    ('GRID',          (0, 0), (-1, -1), 0.3, colors.lightgrey),
                    ('VALIGN',        (0, 0), (-1, -1), 'TOP'),
                    ('TOPPADDING',    (0, 0), (-1, -1), 3),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
                    ('LEFTPADDING',   (0, 0), (-1, -1), 4),
                ]))
                story.append(tbl)
                story.append(PageBreak())

            # ════════════════════════════════════════════════
            #  OTHER SECTIONS
            # ════════════════════════════════════════════════
            for heading, content_list, table_data in self.sections:
                story.append(Paragraph(heading, h1_style))
                story.append(HRFlowable(width='100%', thickness=0.5,
                                        color=colors.lightgrey))
                story.append(Spacer(1, 0.2*cm))
                for item in content_list:
                    s = warn_style if str(item).startswith('⚠') else body_style
                    safe = (str(item)
                            .replace('&', '&amp;')
                            .replace('<', '&lt;')
                            .replace('>', '&gt;'))
                    story.append(Paragraph(safe, s))

                # Optional inline table
                if table_data and len(table_data) > 1:
                    story.append(Spacer(1, 0.2*cm))
                    t = Table(table_data, repeatRows=1)
                    t.setStyle(TableStyle([
                        ('BACKGROUND', (0,0), (-1, 0), ACCENT),
                        ('TEXTCOLOR',  (0,0), (-1, 0), colors.white),
                        ('FONTNAME',   (0,0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE',   (0,0), (-1,-1), 8),
                        ('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.white, LIGHT]),
                        ('GRID',       (0,0), (-1,-1), 0.3, colors.lightgrey),
                        ('TOPPADDING', (0,0), (-1,-1), 4),
                        ('BOTTOMPADDING',(0,0),(-1,-1), 4),
                    ]))
                    story.append(t)

                story.append(Spacer(1, 0.4*cm))

            # ── Footer ────────────────────────────────────────
            story.append(HRFlowable(width='100%', thickness=1, color=ACCENT))
            story.append(Spacer(1, 0.3*cm))
            story.append(Paragraph(
                'This report is generated by ForensiTrace — Forensic Evidence Analysis &amp; '
                'Timeline Reconstruction System | The Islamia University of Bahawalpur | '
                'Spring 2026 | All evidence must be handled per chain-of-custody procedures.',
                ParagraphStyle('footer', parent=styles['Normal'],
                               fontSize=7.5, textColor=colors.grey, alignment=1)
            ))

            doc.build(story)
            return True
        except Exception as e:
            print(f'[ERROR] PDF build failed: {e}')
            import traceback; traceback.print_exc()
            return False


# ═══════════════════════════════════════════════════════════════
#  5.  Browser History & Cache Analyzer
# ═══════════════════════════════════════════════════════════════

import sqlite3
import shutil
import tempfile
import glob

class BrowserHistoryAnalyzer:
    """Extracts visited URLs, downloads, searches from Chrome & Firefox."""

    CHROME_PATHS = [
        os.path.expanduser(r'~\AppData\Local\Google\Chrome\User Data\Default\History'),
        os.path.expanduser('~/.config/google-chrome/Default/History'),
        os.path.expanduser('~/Library/Application Support/Google/Chrome/Default/History'),
    ]
    EDGE_PATHS = [
        os.path.expanduser(r'~\AppData\Local\Microsoft\Edge\User Data\Default\History'),
    ]
    FIREFOX_PATHS = [
        os.path.expanduser(r'~\AppData\Roaming\Mozilla\Firefox\Profiles'),
        os.path.expanduser('~/.mozilla/firefox'),
    ]

    def __init__(self):
        self.results = {
            'chrome' : {'history': [], 'downloads': [], 'searches': []},
            'edge'   : {'history': [], 'downloads': [], 'searches': []},
            'firefox': {'history': [], 'downloads': [], 'searches': []},
        }
        self.errors = []

    # ── Find DB files ─────────────────────────────────────────
    def _find_firefox_db(self):
        for base in self.FIREFOX_PATHS:
            if not os.path.isdir(base):
                continue
            for root, dirs, files in os.walk(base):
                for f in files:
                    if f == 'places.sqlite':
                        return os.path.join(root, f)
        return None

    def _find_chromium_db(self, paths):
        for p in paths:
            if os.path.isfile(p):
                return p
        return None

    # ── Query Chromium-based (Chrome / Edge) ──────────────────
    def _query_chromium(self, db_path: str, label: str):
        tmp = tempfile.mktemp(suffix='.db')
        try:
            shutil.copy2(db_path, tmp)
            conn = sqlite3.connect(tmp)
            cur  = conn.cursor()

            # History
            try:
                cur.execute("""
                    SELECT url, title, visit_count,
                           datetime(last_visit_time/1000000-11644473600,'unixepoch') as last_visit
                    FROM urls ORDER BY last_visit_time DESC LIMIT 500
                """)
                for row in cur.fetchall():
                    self.results[label]['history'].append({
                        'url'        : row[0],
                        'title'      : row[1] or '',
                        'visit_count': row[2],
                        'last_visit' : row[3],
                    })
            except Exception as e:
                self.errors.append(f'{label} history: {e}')

            # Downloads
            try:
                cur.execute("""
                    SELECT target_path, tab_url,
                           datetime(start_time/1000000-11644473600,'unixepoch') as start,
                           received_bytes, state
                    FROM downloads ORDER BY start_time DESC LIMIT 200
                """)
                for row in cur.fetchall():
                    self.results[label]['downloads'].append({
                        'path'  : row[0],
                        'url'   : row[1],
                        'time'  : row[2],
                        'size'  : row[3],
                        'state' : row[4],
                    })
            except Exception as e:
                self.errors.append(f'{label} downloads: {e}')

            # Searches (keyword_search_terms)
            try:
                cur.execute("""
                    SELECT lower_term, url_id FROM keyword_search_terms
                    ORDER BY url_id DESC LIMIT 200
                """)
                for row in cur.fetchall():
                    self.results[label]['searches'].append({
                        'term'  : row[0],
                        'url_id': row[1],
                    })
            except Exception as e:
                self.errors.append(f'{label} searches: {e}')

            conn.close()
        except Exception as e:
            self.errors.append(f'{label} DB error: {e}')
        finally:
            try: os.remove(tmp)
            except: pass

    # ── Query Firefox ─────────────────────────────────────────
    def _query_firefox(self, db_path: str):
        tmp = tempfile.mktemp(suffix='.db')
        try:
            shutil.copy2(db_path, tmp)
            conn = sqlite3.connect(tmp)
            cur  = conn.cursor()

            try:
                cur.execute("""
                    SELECT p.url, p.title, p.visit_count,
                           datetime(h.visit_date/1000000,'unixepoch') as visit_time
                    FROM moz_places p
                    LEFT JOIN moz_historyvisits h ON h.place_id = p.id
                    ORDER BY h.visit_date DESC LIMIT 500
                """)
                for row in cur.fetchall():
                    self.results['firefox']['history'].append({
                        'url'        : row[0],
                        'title'      : row[1] or '',
                        'visit_count': row[2],
                        'last_visit' : row[3],
                    })
            except Exception as e:
                self.errors.append(f'firefox history: {e}')

            # Downloads from moz_annos
            try:
                cur.execute("""
                    SELECT p.url,
                           datetime(a.dateAdded/1000000,'unixepoch') as dl_time
                    FROM moz_annos a
                    JOIN moz_places p ON p.id = a.place_id
                    WHERE a.anno_attribute_id IN (
                        SELECT id FROM moz_anno_attributes
                        WHERE name='downloads/destinationFileName'
                    )
                    ORDER BY a.dateAdded DESC LIMIT 200
                """)
                for row in cur.fetchall():
                    self.results['firefox']['downloads'].append({
                        'url' : row[0], 'time': row[1],
                        'path': '', 'size': 0, 'state': 'N/A',
                    })
            except Exception as e:
                self.errors.append(f'firefox downloads: {e}')

            conn.close()
        except Exception as e:
            self.errors.append(f'firefox DB error: {e}')
        finally:
            try: os.remove(tmp)
            except: pass

    # ── Public method ─────────────────────────────────────────
    def analyse(self) -> dict:
        db = self._find_chromium_db(self.CHROME_PATHS)
        if db: self._query_chromium(db, 'chrome')

        db = self._find_chromium_db(self.EDGE_PATHS)
        if db: self._query_chromium(db, 'edge')

        db = self._find_firefox_db()
        if db: self._query_firefox(db)

        return self.results

    def summary(self) -> dict:
        return {
            'chrome_history'  : len(self.results['chrome']['history']),
            'chrome_downloads': len(self.results['chrome']['downloads']),
            'chrome_searches' : len(self.results['chrome']['searches']),
            'edge_history'    : len(self.results['edge']['history']),
            'edge_downloads'  : len(self.results['edge']['downloads']),
            'firefox_history' : len(self.results['firefox']['history']),
            'firefox_downloads':len(self.results['firefox']['downloads']),
            'errors'          : self.errors,
        }


# ═══════════════════════════════════════════════════════════════
#  6.  Email Header Analyzer
# ═══════════════════════════════════════════════════════════════

import email
import re
import ipaddress

class EmailHeaderAnalyzer:
    """Parse raw email headers to detect spoofing, trace IPs, check SPF/DKIM."""

    PRIVATE_RANGES = [
        ipaddress.ip_network('10.0.0.0/8'),
        ipaddress.ip_network('172.16.0.0/12'),
        ipaddress.ip_network('192.168.0.0/16'),
        ipaddress.ip_network('127.0.0.0/8'),
    ]

    def __init__(self, raw_header: str):
        self.raw    = raw_header
        self.result = {}

    def _is_private(self, ip_str: str) -> bool:
        try:
            ip = ipaddress.ip_address(ip_str)
            return any(ip in net for net in self.PRIVATE_RANGES)
        except ValueError:
            return False

    def _extract_ips(self, text: str) -> list:
        pattern = r'\b(?:\d{1,3}\.){3}\d{1,3}\b'
        found   = re.findall(pattern, text)
        public  = [ip for ip in found if not self._is_private(ip)]
        return list(dict.fromkeys(public))   # unique, order-preserving

    def analyse(self) -> dict:
        r = {}

        # Try Gmail/web info-box format first
        gmail = self._parse_gmail_format(self.raw)
        if gmail:
            r.update(gmail)
        else:
            # Standard RFC email header parsing
            try:
                msg = email.message_from_string(self.raw)
            except Exception as e:
                return {'error': str(e)}
            r['from']        = msg.get('From',           'N/A')
            r['to']          = msg.get('To',             'N/A')
            r['subject']     = msg.get('Subject',        'N/A')
            r['date']        = msg.get('Date',           'N/A')
            r['reply_to']    = msg.get('Reply-To',       'N/A')
            r['return_path'] = msg.get('Return-Path',    'N/A')
            r['message_id']  = msg.get('Message-ID',     'N/A')
            r['x_mailer']    = msg.get('X-Mailer',       'N/A')
            r['x_originating_ip'] = msg.get('X-Originating-IP', 'N/A')
            r['spf']         = msg.get('Received-SPF',   'Not found')
            r['dkim']        = msg.get('DKIM-Signature', 'Not found')
            r['dmarc']       = msg.get('Authentication-Results', 'Not found')
            received         = msg.get_all('Received') or []
            r['received_chain'] = received
            r['hop_count']   = len(received)
            r['public_ips']  = self._extract_ips('\n'.join(received))

        # Defaults
        for k in ('from','to','subject','date','reply_to','return_path',
                  'message_id','x_mailer','x_originating_ip','spf','dkim','dmarc'):
            r.setdefault(k, 'N/A')
        r.setdefault('received_chain', [])
        r.setdefault('hop_count', 0)
        r.setdefault('public_ips', self._extract_ips(self.raw))

        # Spoofing & security checks
        warnings = []
        def _domain(addr):
            m = re.search(r'@([\w.\-]+)', str(addr))
            return m.group(1).lower() if m else ''

        from_dom   = _domain(r.get('from',''))
        reply_dom  = _domain(r.get('reply_to',''))
        return_dom = _domain(r.get('return_path',''))
        spf_val    = str(r.get('spf','')).lower()
        dkim_val   = str(r.get('dkim','')).lower()
        dmarc_val  = str(r.get('dmarc','')).lower()

        if reply_dom and from_dom and reply_dom != from_dom:
            warnings.append(f'⚠ Reply-To ({reply_dom}) ≠ From ({from_dom}) — possible spoofing!')
        if return_dom and from_dom and return_dom != from_dom:
            warnings.append(f'⚠ Return-Path domain ({return_dom}) ≠ From — possible spoofing!')
        if 'fail' in spf_val:
            warnings.append('⚠ SPF: FAILED — sender IP not authorised!')
        elif 'pass' in spf_val:
            warnings.append('✅ SPF: PASS — sender IP is authorised.')
        else:
            warnings.append('ℹ SPF: Not found in headers.')
        if 'pass' in dkim_val:
            warnings.append('✅ DKIM: PASS — email signature verified.')
        elif 'fail' in dkim_val:
            warnings.append('⚠ DKIM: FAILED — signature mismatch!')
        else:
            warnings.append('⚠ DKIM signature not found in headers.')
        if 'pass' in dmarc_val:
            warnings.append('✅ DMARC: PASS')
        elif 'fail' in dmarc_val:
            warnings.append('⚠ DMARC: FAILED')
        if not warnings:
            warnings.append('✅ No obvious spoofing indicators detected.')

        r['warnings'] = warnings
        self.result   = r
        return r

    def _parse_gmail_format(self, text):
        # Parse Gmail info-box style (non RFC) format
        r = {}
        patterns = {
            'message_id': r'Message ID[<\s:]+([^\n>]+)',
            'date'       : r'Created at[:\s]+([^\n]+)',
            'from'       : r'From[:\s]+([^\n]+)',
            'to'         : r'To[:\s]+([^\n]+)',
            'subject'    : r'Subject[:\s]+([^\n]+)',
            'spf'        : r'SPF[:\s]+([^\n]+)',
            'dkim'       : r'DKIM[:\s]+([^\n]+)',
            'dmarc'      : r'DMARC[:\s]+([^\n]+)',
        }
        found = 0
        for key, pattern in patterns.items():
            m = re.search(pattern, text, re.IGNORECASE)
            if m:
                r[key] = m.group(1).strip().rstrip('>')
                found += 1
        if found < 2:
            return None
        # Extract IP from SPF line
        ip_m = re.search(r'SPF.*?with IP\s+([\d.]+)', text, re.IGNORECASE)
        if ip_m:
            r['public_ips'] = [ip_m.group(1)]
            r['x_originating_ip'] = ip_m.group(1)
        r.setdefault('reply_to', 'N/A')
        r.setdefault('return_path', 'N/A')
        r.setdefault('x_mailer', 'N/A')
        r.setdefault('received_chain', [])
        r.setdefault('hop_count', 0)
        r.setdefault('public_ips', [])
        return r

# ═══════════════════════════════════════════════════════════════
#  7.  Windows Event Log Parser
# ═══════════════════════════════════════════════════════════════

class WindowsEventLogParser:
    """
    Parse Windows .evtx files or read live Security/System logs.
    Flags: login attempts, failed passwords, USB events.
    """

    INTERESTING_EVENT_IDS = {
        4624: 'Successful Logon',
        4625: 'Failed Logon',
        4634: 'Logoff',
        4648: 'Logon with explicit credentials',
        4672: 'Special privileges assigned',
        4688: 'Process Created',
        4698: 'Scheduled Task Created',
        4720: 'User Account Created',
        4726: 'User Account Deleted',
        4732: 'User added to privileged group',
        4776: 'Credential Validation',
        6416: 'New external device recognised (USB)',
        7045: 'New Service Installed',
    }

    def __init__(self):
        self.events = []
        self.errors = []

    # ── Parse .evtx file ─────────────────────────────────────
    def parse_evtx(self, path: str) -> list:
        try:
            import Evtx.Evtx as evtx
            import xml.etree.ElementTree as ET

            events = []
            with evtx.Evtx(path) as log:
                for record in log.records():
                    try:
                        xml_str = record.xml()
                        root    = ET.fromstring(xml_str)
                        ns      = {'e': 'http://schemas.microsoft.com/win/2004/08/events/event'}

                        eid_el = root.find('.//e:EventID', ns)
                        tim_el = root.find('.//e:TimeCreated', ns)
                        cmp_el = root.find('.//e:Computer', ns)

                        eid  = int(eid_el.text) if eid_el is not None else 0
                        time = tim_el.attrib.get('SystemTime','') if tim_el is not None else ''
                        comp = cmp_el.text if cmp_el is not None else ''

                        label = self.INTERESTING_EVENT_IDS.get(eid, '')
                        evt   = {
                            'event_id'   : eid,
                            'label'      : label,
                            'time'       : time[:19].replace('T',' '),
                            'computer'   : comp,
                            'raw_xml'    : xml_str[:500],
                            'interesting': bool(label),
                        }

                        # Extract user / IP from EventData
                        for data in root.findall('.//e:Data', ns):
                            name = data.attrib.get('Name','')
                            val  = data.text or ''
                            if name in ('SubjectUserName','TargetUserName','IpAddress',
                                        'ProcessName','ServiceName','DeviceName'):
                                evt[name] = val

                        events.append(evt)
                    except Exception:
                        continue

            self.events = events
            return events

        except ImportError:
            self.errors.append('python-evtx not installed. Run: pip install python-evtx')
            return []
        except Exception as e:
            self.errors.append(f'EVTX parse error: {e}')
            return []

    # ── Live Windows logs (win32evtlog) ──────────────────────
    def read_live(self, log_name: str = 'Security', count: int = 500) -> list:
        if platform.system() != 'Windows':
            self.errors.append('Live log reading only supported on Windows.')
            return []
        try:
            import win32evtlog
            hand   = win32evtlog.OpenEventLog(None, log_name)
            flags  = win32evtlog.EVENTLOG_BACKWARDS_READ | win32evtlog.EVENTLOG_SEQUENTIAL_READ
            events = []
            total  = 0
            while total < count:
                batch = win32evtlog.ReadEventLog(hand, flags, 0)
                if not batch:
                    break
                for e in batch:
                    eid   = e.EventID & 0xFFFF
                    label = self.INTERESTING_EVENT_IDS.get(eid, '')
                    events.append({
                        'event_id'   : eid,
                        'label'      : label,
                        'time'       : str(e.TimeGenerated),
                        'source'     : e.SourceName,
                        'computer'   : e.ComputerName,
                        'interesting': bool(label),
                        'message'    : str(e.StringInserts or ''),
                    })
                    total += 1
                if total >= count:
                    break
            win32evtlog.CloseEventLog(hand)
            self.events = events
            return events
        except Exception as e:
            self.errors.append(f'Live log error: {e}')
            return []

    def interesting_only(self) -> list:
        return [e for e in self.events if e.get('interesting')]

    def summary(self) -> dict:
        interesting = self.interesting_only()
        counts = {}
        for e in interesting:
            lbl = e.get('label','Unknown')
            counts[lbl] = counts.get(lbl, 0) + 1
        return {
            'total_events'      : len(self.events),
            'interesting_events': len(interesting),
            'event_counts'      : counts,
            'errors'            : self.errors,
        }


# ═══════════════════════════════════════════════════════════════
#  8.  Steganography Detector
# ═══════════════════════════════════════════════════════════════

class SteganographyDetector:
    """
    Detects hidden data in images using:
    - LSB (Least Significant Bit) analysis
    - File size vs content ratio check
    - Appended data after image EOF
    - Chi-square randomness test on LSBs
    """

    def __init__(self, image_path: str):
        self.path   = image_path
        self.result = {}

    def analyse(self) -> dict:
        r = {
            'file'          : self.path,
            'warnings'      : [],
            'lsb_entropy'   : None,
            'appended_data' : False,
            'appended_bytes': 0,
            'chi_square'    : None,
            'verdict'       : 'Clean',
        }

        if not os.path.isfile(self.path):
            r['warnings'].append('File not found.')
            self.result = r
            return r

        ext = Path(self.path).suffix.lower()

        # ── Appended data check ───────────────────────────────
        try:
            with open(self.path, 'rb') as f:
                data = f.read()

            # JPEG EOF marker
            if ext in ('.jpg', '.jpeg'):
                eof_pos = data.rfind(b'\xff\xd9')
                if eof_pos != -1 and eof_pos < len(data) - 2:
                    extra = len(data) - eof_pos - 2
                    r['appended_data']  = True
                    r['appended_bytes'] = extra
                    r['warnings'].append(f'⚠ {extra} bytes of data found AFTER JPEG EOF marker — possible hidden payload!')

            # PNG IEND check
            if ext == '.png':
                iend_pos = data.rfind(b'IEND')
                if iend_pos != -1:
                    after = len(data) - iend_pos - 8
                    if after > 0:
                        r['appended_data']  = True
                        r['appended_bytes'] = after
                        r['warnings'].append(f'⚠ {after} bytes found AFTER PNG IEND chunk — possible hidden data!')
        except Exception as e:
            r['warnings'].append(f'Appended data check error: {e}')

        # ── LSB entropy + Chi-square via Pillow ───────────────
        try:
            from PIL import Image
            import math, struct

            img  = Image.open(self.path).convert('RGB')
            pix  = list(img.getdata())

            # Extract LSBs from red channel
            lsbs = [p[0] & 1 for p in pix[:50000]]

            # Chi-square test (expected ~50% 0s and 50% 1s if stego)
            ones  = sum(lsbs)
            zeros = len(lsbs) - ones
            n     = len(lsbs)
            exp   = n / 2
            chi   = ((ones - exp)**2 + (zeros - exp)**2) / exp if exp else 0
            r['chi_square'] = round(chi, 4)

            # LSB entropy
            p1 = ones / n if n else 0
            p0 = 1 - p1
            def ent(p):
                return -p * math.log2(p) if p > 0 else 0
            entropy = ent(p1) + ent(p0)
            r['lsb_entropy'] = round(entropy, 6)

            # Stego detection: JPEG camera photos naturally have near-uniform LSBs
            # due to compression, so we use stricter thresholds
            # Real stego tools produce chi < 1 and entropy > 0.999
            if chi < 1.0 and entropy > 0.999 and not r['appended_data']:
                r['warnings'].append('⚠ LSB distribution extremely uniform — possible LSB steganography.')
                r['verdict'] = 'SUSPICIOUS — Possible Steganography'
            elif r['appended_data']:
                r['warnings'].append('⚠ Extra data found after image EOF — investigate further.')
                r['verdict'] = 'SUSPICIOUS — Appended Data Found'
            else:
                r['warnings'].append('✅ LSB distribution normal for a regular camera/digital image.')
                r['verdict'] = 'Clean'

        except ImportError:
            r['warnings'].append('Pillow not installed — install with: pip install pillow')
        except Exception as e:
            r['warnings'].append(f'LSB analysis error: {e}')

        if not r['warnings']:
            r['warnings'].append('✅ No hidden data detected.')

        self.result = r
        return r


# ═══════════════════════════════════════════════════════════════
#  9.  USB Activity Monitor
# ═══════════════════════════════════════════════════════════════

class USBActivityMonitor:
    """
    Detects USB device history from:
    - Windows Registry (SYSTEM hive via offline parsing)
    - Windows Event Logs (Event ID 6416 / 2003)
    - /var/log/syslog or dmesg on Linux
    """

    def __init__(self):
        self.devices = []
        self.errors  = []

    # ── Windows registry via winreg ───────────────────────────
    def scan_windows_registry(self) -> list:
        if platform.system() != 'Windows':
            self.errors.append('Registry scan only works on Windows.')
            return []
        try:
            import winreg
            key_path = r'SYSTEM\CurrentControlSet\Enum\USBSTOR'
            root     = winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE, key_path)
            devices  = []
            i = 0
            while True:
                try:
                    subkey_name = winreg.EnumKey(root, i)
                    subkey      = winreg.OpenKey(root, subkey_name)
                    j = 0
                    while True:
                        try:
                            instance = winreg.EnumKey(subkey, j)
                            inst_key = winreg.OpenKey(subkey, instance)
                            try:
                                friendly, _ = winreg.QueryValueEx(inst_key, 'FriendlyName')
                            except:
                                friendly = subkey_name
                            try:
                                serial = instance.split('&')[0]
                            except:
                                serial = instance

                            devices.append({
                                'device_type' : subkey_name,
                                'serial'      : serial,
                                'friendly_name': friendly,
                                'instance'    : instance,
                                'source'      : 'Registry',
                            })
                            j += 1
                        except OSError:
                            break
                    i += 1
                except OSError:
                    break
            self.devices = devices
            return devices
        except Exception as e:
            self.errors.append(f'Registry scan error: {e}')
            return []

    # ── Linux syslog / dmesg scan ─────────────────────────────
    def scan_linux_logs(self) -> list:
        devices = []
        log_files = ['/var/log/syslog', '/var/log/kern.log', '/var/log/messages']
        for lf in log_files:
            if not os.path.isfile(lf):
                continue
            try:
                with open(lf, 'r', errors='replace') as f:
                    for line in f:
                        if 'usb' in line.lower() and any(
                            kw in line.lower() for kw in
                            ['new usb device', 'usb-storage', 'scsi', 'product:', 'manufacturer:']):
                            devices.append({
                                'log_file' : lf,
                                'line'     : line.strip()[:200],
                                'source'   : 'syslog',
                            })
            except Exception as e:
                self.errors.append(f'Log scan error ({lf}): {e}')
        self.devices = devices
        return devices

    def scan(self) -> list:
        if platform.system() == 'Windows':
            return self.scan_windows_registry()
        else:
            return self.scan_linux_logs()

    def summary(self) -> dict:
        return {
            'total_devices': len(self.devices),
            'devices'      : self.devices,
            'errors'       : self.errors,
        }


# ═══════════════════════════════════════════════════════════════
#  10. Memory Dump Analyzer
# ═══════════════════════════════════════════════════════════════

class MemoryDumpAnalyzer:
    """
    Basic forensic analysis of a raw memory dump file.
    Extracts: strings, URLs, IPs, process hints, Windows artifacts.
    No Volatility dependency — pure Python.
    """

    URL_RE     = re.compile(rb'https?://[^\x00-\x1f\x7f-\xff ]{8,200}')
    IP_RE      = re.compile(rb'\b(?:\d{1,3}\.){3}\d{1,3}\b')
    EMAIL_RE   = re.compile(rb'[\w.\-]+@[\w.\-]+\.\w{2,6}')
    WIN_PROC   = [b'svchost', b'lsass', b'csrss', b'winlogon', b'explorer',
                  b'cmd.exe', b'powershell', b'rundll32', b'taskhost', b'conhost']
    SUSPICIOUS = [b'mimikatz', b'metasploit', b'meterpreter', b'nc.exe',
                  b'ncat', b'netcat', b'cobaltstrike', b'empire', b'shellcode']

    def __init__(self, dump_path: str, chunk_size: int = 10 * 1024 * 1024):
        self.dump_path  = dump_path
        self.chunk_size = chunk_size
        self.result     = {}

    def analyse(self) -> dict:
        r = {
            'file'           : self.dump_path,
            'size_human'     : 'N/A',
            'urls'           : [],
            'ips'            : [],
            'emails'         : [],
            'processes_found': [],
            'suspicious_strings': [],
            'warnings'       : [],
        }

        if not os.path.isfile(self.dump_path):
            r['warnings'].append('Dump file not found.')
            self.result = r
            return r

        size = os.path.getsize(self.dump_path)
        r['size_human'] = FileMetadataExtractor._human(size)

        urls_set  = set()
        ips_set   = set()
        email_set = set()
        procs     = set()
        sus_set   = set()

        try:
            with open(self.dump_path, 'rb') as f:
                while True:
                    chunk = f.read(self.chunk_size)
                    if not chunk:
                        break
                    for m in self.URL_RE.findall(chunk):
                        urls_set.add(m.decode('latin-1', errors='replace')[:200])
                    for m in self.IP_RE.findall(chunk):
                        ips_set.add(m.decode('latin-1'))
                    for m in self.EMAIL_RE.findall(chunk):
                        email_set.add(m.decode('latin-1', errors='replace'))
                    for p in self.WIN_PROC:
                        if p in chunk.lower():
                            procs.add(p.decode())
                    for s in self.SUSPICIOUS:
                        if s in chunk.lower():
                            sus_set.add(s.decode())

            r['urls']               = sorted(urls_set)[:200]
            r['ips']                = sorted(ips_set)[:200]
            r['emails']             = sorted(email_set)[:100]
            r['processes_found']    = sorted(procs)
            r['suspicious_strings'] = sorted(sus_set)

            if sus_set:
                r['warnings'].append(
                    f'⚠ SUSPICIOUS strings found: {", ".join(sus_set)} — possible malware!')
            if not r['warnings']:
                r['warnings'].append('✅ No obvious malware strings detected.')

        except Exception as e:
            r['warnings'].append(f'Analysis error: {e}')

        self.result = r
        return r


# ═══════════════════════════════════════════════════════════════
#  11. Deleted File Recovery (Filesystem Artifact Scanner)
# ═══════════════════════════════════════════════════════════════

class DeletedFileScanner:
    """
    Scans a disk image or folder for deleted file artifacts:
    - File carving by magic bytes from raw image
    - Recycle Bin ($Recycle.Bin) analysis on Windows
    - Temp folder remnants
    - Recently deleted (Trash) on Linux/Mac
    """

    CARVE_SIGS = {
        b'\xff\xd8\xff'    : ('jpg',  b'\xff\xd9',    5*1024*1024),
        b'\x89PNG\r\n\x1a\n': ('png', b'IEND\xaeB`\x82', 5*1024*1024),
        b'%PDF'            : ('pdf',  b'%%EOF',       20*1024*1024),
        b'PK\x03\x04'      : ('zip',  b'PK\x05\x06', 50*1024*1024),
        b'MZ'              : ('exe',  None,           2*1024*1024),
    }

    def __init__(self):
        self.found   = []
        self.errors  = []

    # ── Recycle Bin ───────────────────────────────────────────
    def scan_recycle_bin(self) -> list:
        items = []
        if platform.system() != 'Windows':
            self.errors.append('Recycle Bin scan only available on Windows.')
            return items
        for drive in 'CDEFGH':
            rb = f'{drive}:\\$Recycle.Bin'
            if not os.path.isdir(rb):
                continue
            try:
                for root, dirs, files in os.walk(rb):
                    for fn in files:
                        fp   = os.path.join(root, fn)
                        size = 0
                        try: size = os.path.getsize(fp)
                        except: pass
                        items.append({
                            'source'    : 'Recycle Bin',
                            'path'      : fp,
                            'name'      : fn,
                            'size_bytes': size,
                            'size_human': FileMetadataExtractor._human(size),
                        })
            except Exception as e:
                self.errors.append(f'Recycle Bin ({drive}:) error: {e}')
        return items

    # ── Temp folder ───────────────────────────────────────────
    def scan_temp_folder(self) -> list:
        items   = []
        folders = [tempfile.gettempdir()]
        if platform.system() == 'Windows':
            folders += [
                os.path.expanduser(r'~\AppData\Local\Temp'),
                r'C:\Windows\Temp',
            ]
        for folder in folders:
            if not os.path.isdir(folder):
                continue
            try:
                for fn in os.listdir(folder):
                    fp   = os.path.join(folder, fn)
                    size = 0
                    try: size = os.path.getsize(fp)
                    except: pass
                    items.append({
                        'source'    : 'Temp Folder',
                        'path'      : fp,
                        'name'      : fn,
                        'size_bytes': size,
                        'size_human': FileMetadataExtractor._human(size),
                    })
            except Exception as e:
                self.errors.append(f'Temp scan error ({folder}): {e}')
        return items

    # ── Linux Trash ───────────────────────────────────────────
    def scan_linux_trash(self) -> list:
        items   = []
        trashes = [
            os.path.expanduser('~/.local/share/Trash/files'),
            '/root/.local/share/Trash/files',
        ]
        for t in trashes:
            if not os.path.isdir(t):
                continue
            try:
                for fn in os.listdir(t):
                    fp   = os.path.join(t, fn)
                    size = 0
                    try: size = os.path.getsize(fp)
                    except: pass
                    items.append({
                        'source'    : 'Trash',
                        'path'      : fp,
                        'name'      : fn,
                        'size_bytes': size,
                        'size_human': FileMetadataExtractor._human(size),
                    })
            except Exception as e:
                self.errors.append(f'Trash scan error: {e}')
        return items

    # ── File carving from raw disk image ──────────────────────
    def carve_image(self, image_path: str, output_dir: str,
                    max_files: int = 50) -> list:
        carved = []
        if not os.path.isfile(image_path):
            self.errors.append(f'Image not found: {image_path}')
            return carved
        os.makedirs(output_dir, exist_ok=True)
        try:
            with open(image_path, 'rb') as f:
                data = f.read()
            count = 0
            for sig, (ext, end_sig, max_size) in self.CARVE_SIGS.items():
                pos = 0
                while count < max_files:
                    idx = data.find(sig, pos)
                    if idx == -1:
                        break
                    end = -1
                    if end_sig:
                        end = data.find(end_sig, idx + len(sig))
                        if end == -1 or (end - idx) > max_size:
                            pos = idx + 1
                            continue
                        end += len(end_sig)
                    else:
                        end = min(idx + max_size, len(data))

                    chunk     = data[idx:end]
                    out_name  = f'carved_{count:04d}.{ext}'
                    out_path  = os.path.join(output_dir, out_name)
                    with open(out_path, 'wb') as out:
                        out.write(chunk)
                    carved.append({
                        'source'    : 'File Carving',
                        'offset'    : hex(idx),
                        'name'      : out_name,
                        'path'      : out_path,
                        'size_human': FileMetadataExtractor._human(len(chunk)),
                        'type'      : ext,
                    })
                    count += 1
                    pos    = idx + 1
        except Exception as e:
            self.errors.append(f'Carving error: {e}')
        return carved

    # ── Full scan ─────────────────────────────────────────────
    def scan_all(self) -> list:
        found = []
        found += self.scan_recycle_bin()
        found += self.scan_temp_folder()
        found += self.scan_linux_trash()
        self.found = found
        return found

    def summary(self) -> dict:
        return {
            'total_artifacts': len(self.found),
            'artifacts'      : self.found,
            'errors'         : self.errors,
        }


# ═══════════════════════════════════════════════════════════════
#  12. Password Strength Checker
# ═══════════════════════════════════════════════════════════════

class PasswordStrengthChecker:
    """
    Forensic password strength analysis:
    - Length, entropy, character diversity
    - Pattern detection (dictionary words, keyboard walks, dates)
    - Crack time estimation
    - Strength score 0-100
    """

    COMMON_PASSWORDS = {
        'password','123456','password1','qwerty','abc123','letmein',
        'monkey','1234567890','iloveyou','admin','welcome','login',
        'master','hello','shadow','sunshine','princess','dragon',
        'pass','test','123123','111111','000000','superman','batman'
    }

    KEYBOARD_WALKS = ['qwerty','asdfgh','zxcvbn','qweasd','1234567890',
                      '!@#$%^','qwertyuiop','asdfghjkl']

    def analyse(self, password: str) -> dict:
        import math, string, re

        r = {
            'password'        : '*' * len(password),  # masked
            'length'          : len(password),
            'entropy_bits'    : 0.0,
            'score'           : 0,
            'strength_label'  : '',
            'crack_time'      : '',
            'has_upper'       : False,
            'has_lower'       : False,
            'has_digit'       : False,
            'has_special'     : False,
            'is_common'       : False,
            'has_keyboard_walk': False,
            'has_date_pattern': False,
            'has_repeated'    : False,
            'suggestions'     : [],
            'warnings'        : [],
        }

        if not password:
            r['strength_label'] = 'Empty'
            return r

        # Character set size
        charset = 0
        if any(c.isupper() for c in password):   r['has_upper']   = True; charset += 26
        if any(c.islower() for c in password):   r['has_lower']   = True; charset += 26
        if any(c.isdigit() for c in password):   r['has_digit']   = True; charset += 10
        if any(c in string.punctuation for c in password): r['has_special'] = True; charset += 32

        # Entropy
        if charset > 0:
            r['entropy_bits'] = round(len(password) * math.log2(charset), 2)

        # Common password check
        if password.lower() in self.COMMON_PASSWORDS:
            r['is_common'] = True
            r['warnings'].append('⚠ This is a very common password — instantly crackable!')

        # Keyboard walk
        pw_lower = password.lower()
        for walk in self.KEYBOARD_WALKS:
            if walk in pw_lower:
                r['has_keyboard_walk'] = True
                r['warnings'].append(f'⚠ Keyboard pattern detected: "{walk}"')
                break

        # Date pattern
        if re.search(r'\b(19|20)\d{2}\b|\b\d{2}[/\-]\d{2}[/\-]\d{2,4}\b', password):
            r['has_date_pattern'] = True
            r['warnings'].append('⚠ Date pattern detected — easily guessable.')

        # Repeated chars
        if re.search(r'(.)\1{2,}', password):
            r['has_repeated'] = True
            r['warnings'].append('⚠ Repeated characters detected.')

        # Score (0-100)
        score = 0
        score += min(30, len(password) * 2)           # length (max 30)
        score += 10 if r['has_upper']   else 0
        score += 10 if r['has_lower']   else 0
        score += 10 if r['has_digit']   else 0
        score += 15 if r['has_special'] else 0
        score += min(25, int(r['entropy_bits'] / 3))   # entropy (max 25)
        score -= 30 if r['is_common']          else 0
        score -= 15 if r['has_keyboard_walk']  else 0
        score -= 10 if r['has_date_pattern']   else 0
        score -= 10 if r['has_repeated']       else 0
        score  = max(0, min(100, score))
        r['score'] = score

        # Strength label
        if   score >= 80: r['strength_label'] = 'Very Strong 💪'
        elif score >= 60: r['strength_label'] = 'Strong ✅'
        elif score >= 40: r['strength_label'] = 'Moderate ⚠'
        elif score >= 20: r['strength_label'] = 'Weak ❌'
        else:             r['strength_label'] = 'Very Weak 🚨'

        # Crack time estimate (assuming 10B guesses/sec — GPU cracking)
        charset_size = charset if charset > 0 else 1
        combinations = charset_size ** len(password)
        seconds = combinations / 10_000_000_000
        if   seconds < 1:           r['crack_time'] = 'Instantly'
        elif seconds < 60:          r['crack_time'] = f'{int(seconds)} seconds'
        elif seconds < 3600:        r['crack_time'] = f'{int(seconds/60)} minutes'
        elif seconds < 86400:       r['crack_time'] = f'{int(seconds/3600)} hours'
        elif seconds < 2592000:     r['crack_time'] = f'{int(seconds/86400)} days'
        elif seconds < 31536000:    r['crack_time'] = f'{int(seconds/2592000)} months'
        elif seconds < 3153600000:  r['crack_time'] = f'{int(seconds/31536000)} years'
        else:                       r['crack_time'] = 'Centuries+'

        # Suggestions
        if not r['has_upper']:   r['suggestions'].append('Add uppercase letters (A-Z)')
        if not r['has_lower']:   r['suggestions'].append('Add lowercase letters (a-z)')
        if not r['has_digit']:   r['suggestions'].append('Add numbers (0-9)')
        if not r['has_special']: r['suggestions'].append('Add special characters (!@#$%)')
        if len(password) < 12:  r['suggestions'].append('Use at least 12 characters')
        if len(password) < 16:  r['suggestions'].append('16+ characters is ideal for high security')

        return r


# ═══════════════════════════════════════════════════════════════
#  13. Network Packet Analyzer  (.pcap)
# ═══════════════════════════════════════════════════════════════

class NetworkPacketAnalyzer:
    """
    Analyse .pcap files without Wireshark.
    Uses raw struct parsing — no scapy dependency required.
    Extracts: IPs, ports, protocols, suspicious traffic.
    """

    SUSPICIOUS_PORTS = {
        4444: 'Metasploit default',  23: 'Telnet (unencrypted)',
        21: 'FTP (unencrypted)',     69: 'TFTP',
        135: 'RPC',                  137: 'NetBIOS',
        139: 'NetBIOS Session',      445: 'SMB (WannaCry target)',
        1433: 'MSSQL',               3306: 'MySQL exposed',
        3389: 'RDP (brute-force target)', 5900: 'VNC',
        6667: 'IRC (botnet C2)',      8080: 'Alt HTTP',
        31337: 'Elite/backdoor',
    }

    PCAP_MAGIC = b'\xd4\xc3\xb2\xa1'   # little-endian pcap magic
    PCAP_MAGIC_BE = b'\xa1\xb2\xc3\xd4'

    def __init__(self, pcap_path: str):
        self.pcap_path = pcap_path
        self.packets   = []
        self.errors    = []

    def _parse_ip(self, raw: bytes) -> str:
        return '.'.join(str(b) for b in raw)

    def analyse(self) -> dict:
        r = {
            'file'           : self.pcap_path,
            'total_packets'  : 0,
            'connections'    : [],
            'top_talkers'    : {},
            'suspicious'     : [],
            'protocols'      : {},
            'warnings'       : [],
        }

        if not os.path.isfile(self.pcap_path):
            r['warnings'].append('PCAP file not found.')
            return r

        try:
            with open(self.pcap_path, 'rb') as f:
                header = f.read(24)
                if len(header) < 24:
                    r['warnings'].append('File too small to be a valid PCAP.')
                    return r

                magic = header[:4]
                if magic not in (self.PCAP_MAGIC, self.PCAP_MAGIC_BE):
                    r['warnings'].append('Not a valid PCAP file (wrong magic bytes).')
                    return r

                endian = '<' if magic == self.PCAP_MAGIC else '>'
                # network (data link type)
                link_type = struct.unpack(endian + 'I', header[20:24])[0]

                packet_count = 0
                src_counts = {}

                while packet_count < 10000:  # max 10k packets
                    rec_hdr = f.read(16)
                    if len(rec_hdr) < 16:
                        break
                    ts_sec, ts_usec, incl_len, orig_len = struct.unpack(endian+'IIII', rec_hdr)
                    raw = f.read(incl_len)
                    if len(raw) < incl_len:
                        break

                    packet_count += 1
                    pkt = {'ts': ts_sec, 'len': orig_len,
                           'src_ip':'','dst_ip':'','src_port':0,'dst_port':0,'proto':''}

                    # Ethernet frame (link_type=1)
                    eth_offset = 14 if link_type == 1 else 0
                    if len(raw) < eth_offset + 20:
                        continue

                    ip_raw = raw[eth_offset:]
                    version = (ip_raw[0] >> 4)
                    if version != 4:
                        continue  # skip IPv6 for now

                    ihl = (ip_raw[0] & 0x0f) * 4
                    proto_num = ip_raw[9]
                    src_ip = self._parse_ip(ip_raw[12:16])
                    dst_ip = self._parse_ip(ip_raw[16:20])
                    pkt['src_ip'] = src_ip
                    pkt['dst_ip'] = dst_ip

                    proto_map = {6:'TCP', 17:'UDP', 1:'ICMP'}
                    proto = proto_map.get(proto_num, str(proto_num))
                    pkt['proto'] = proto
                    r['protocols'][proto] = r['protocols'].get(proto, 0) + 1

                    src_counts[src_ip] = src_counts.get(src_ip, 0) + 1

                    # TCP/UDP ports
                    if proto_num in (6, 17) and len(ip_raw) >= ihl + 4:
                        transport = ip_raw[ihl:]
                        src_port = struct.unpack('>H', transport[0:2])[0]
                        dst_port = struct.unpack('>H', transport[2:4])[0]
                        pkt['src_port'] = src_port
                        pkt['dst_port'] = dst_port

                        conn = f'{src_ip}:{src_port} → {dst_ip}:{dst_port} [{proto}]'
                        if conn not in r['connections']:
                            r['connections'].append(conn)

                        # Suspicious port check
                        for port in (src_port, dst_port):
                            if port in self.SUSPICIOUS_PORTS:
                                entry = (f'⚠ {conn}  —  Port {port}: '
                                         f'{self.SUSPICIOUS_PORTS[port]}')
                                if entry not in r['suspicious']:
                                    r['suspicious'].append(entry)

                    self.packets.append(pkt)

            r['total_packets'] = packet_count
            r['connections']   = r['connections'][:200]
            # Top talkers
            r['top_talkers'] = dict(sorted(src_counts.items(),
                                            key=lambda x: x[1], reverse=True)[:15])
            if not r['warnings']:
                r['warnings'].append(f'✅ Analysed {packet_count} packets successfully.')

        except Exception as e:
            r['warnings'].append(f'Parse error: {e}')

        return r


# ═══════════════════════════════════════════════════════════════
#  14. File Integrity Monitor
# ═══════════════════════════════════════════════════════════════

class FileIntegrityMonitor:
    """
    Watch a folder and detect any file changes, additions, or deletions.
    Baseline snapshot saved to JSON; subsequent scans compared against it.
    """

    def __init__(self, folder: str, baseline_path: str):
        self.folder        = folder
        self.baseline_path = baseline_path
        self.baseline      = {}
        self.current       = {}
        self.changes       = []

    def _snapshot(self) -> dict:
        snap = {}
        try:
            for root, dirs, files in os.walk(self.folder):
                for fn in files:
                    fp = os.path.join(root, fn)
                    try:
                        st = os.stat(fp)
                        # Quick hash of first 64KB for change detection
                        h = hashlib.md5()
                        with open(fp, 'rb') as f:
                            h.update(f.read(65536))
                        snap[fp] = {
                            'size'    : st.st_size,
                            'modified': st.st_mtime,
                            'md5_head': h.hexdigest(),
                        }
                    except Exception:
                        pass
        except Exception as e:
            pass
        return snap

    def save_baseline(self) -> dict:
        self.baseline = self._snapshot()
        with open(self.baseline_path, 'w') as f:
            json.dump(self.baseline, f, indent=2)
        return {'saved': len(self.baseline), 'path': self.baseline_path}

    def load_baseline(self) -> bool:
        if not os.path.isfile(self.baseline_path):
            return False
        try:
            with open(self.baseline_path, 'r') as f:
                self.baseline = json.load(f)
            return True
        except Exception:
            return False

    def scan(self) -> list:
        self.current = self._snapshot()
        self.changes = []

        # Added files
        for fp in self.current:
            if fp not in self.baseline:
                self.changes.append({
                    'type': 'ADDED', 'file': fp,
                    'detail': f'New file detected — size: {FileMetadataExtractor._human(self.current[fp]["size"])}'
                })

        # Deleted files
        for fp in self.baseline:
            if fp not in self.current:
                self.changes.append({'type': 'DELETED', 'file': fp, 'detail': 'File removed!'})

        # Modified files
        for fp in self.current:
            if fp in self.baseline:
                old = self.baseline[fp]
                new = self.current[fp]
                if old['md5_head'] != new['md5_head']:
                    self.changes.append({
                        'type'  : 'MODIFIED',
                        'file'  : fp,
                        'detail': (f'Content changed — '
                                   f'size: {FileMetadataExtractor._human(old["size"])} → '
                                   f'{FileMetadataExtractor._human(new["size"])}')
                    })

        return self.changes

    def summary(self) -> dict:
        added    = [c for c in self.changes if c['type'] == 'ADDED']
        deleted  = [c for c in self.changes if c['type'] == 'DELETED']
        modified = [c for c in self.changes if c['type'] == 'MODIFIED']
        return {
            'folder'         : self.folder,
            'baseline_files' : len(self.baseline),
            'current_files'  : len(self.current),
            'total_changes'  : len(self.changes),
            'added'          : len(added),
            'deleted'        : len(deleted),
            'modified'       : len(modified),
            'changes'        : self.changes,
        }


# ═══════════════════════════════════════════════════════════════
#  15. Malware Hash Lookup (VirusTotal)
# ═══════════════════════════════════════════════════════════════

class MalwareHashLookup:
    """
    Lookup file hash against VirusTotal API v3.
    Returns detection ratio, engine verdicts, file info.
    User must supply their own free VT API key.
    """

    VT_URL = 'https://www.virustotal.com/api/v3/files/{}'

    def __init__(self, api_key: str = ''):
        self.api_key = api_key

    def lookup(self, file_hash: str) -> dict:
        r = {
            'hash'            : file_hash,
            'found'           : False,
            'malicious'       : 0,
            'suspicious'      : 0,
            'harmless'        : 0,
            'undetected'      : 0,
            'total_engines'   : 0,
            'detection_ratio' : '',
            'file_name'       : '',
            'file_type'       : '',
            'file_size'       : '',
            'first_submission': '',
            'last_analysis'   : '',
            'verdicts'        : [],
            'verdict_label'   : '',
            'error'           : '',
        }

        if not self.api_key:
            r['error'] = 'No API key provided. Get a free key at virustotal.com'
            return r

        try:
            import urllib.request
            url = self.VT_URL.format(file_hash.lower().strip())
            req = urllib.request.Request(url, headers={'x-apikey': self.api_key})
            with urllib.request.urlopen(req, timeout=10) as resp:
                data = json.loads(resp.read())

            attrs = data.get('data', {}).get('attributes', {})
            stats = attrs.get('last_analysis_stats', {})

            r['found']        = True
            r['malicious']    = stats.get('malicious', 0)
            r['suspicious']   = stats.get('suspicious', 0)
            r['harmless']     = stats.get('harmless', 0)
            r['undetected']   = stats.get('undetected', 0)
            total = sum(stats.values())
            r['total_engines']= total
            r['detection_ratio'] = f'{r["malicious"]}/{total}'

            names = attrs.get('names', [])
            r['file_name']    = names[0] if names else 'N/A'
            r['file_type']    = attrs.get('type_description', 'N/A')
            r['file_size']    = FileMetadataExtractor._human(attrs.get('size', 0))
            r['first_submission'] = attrs.get('first_submission_date', 'N/A')
            r['last_analysis']    = attrs.get('last_analysis_date', 'N/A')

            # Top engine verdicts
            engines = attrs.get('last_analysis_results', {})
            for eng, result in list(engines.items())[:20]:
                cat = result.get('category','')
                if cat in ('malicious','suspicious'):
                    r['verdicts'].append(f'{eng}: {result.get("result","?")}')

            if   r['malicious'] >= 5:  r['verdict_label'] = '🚨 MALWARE DETECTED'
            elif r['malicious'] > 0:   r['verdict_label'] = '⚠ Possibly Malicious'
            elif r['suspicious'] > 0:  r['verdict_label'] = '⚡ Suspicious'
            else:                      r['verdict_label'] = '✅ Clean'

        except Exception as e:
            err = str(e)
            if '404' in err:
                r['error'] = 'Hash not found in VirusTotal database.'
            elif '401' in err:
                r['error'] = 'Invalid API key.'
            elif '429' in err:
                r['error'] = 'Rate limit exceeded. Try again in 1 minute.'
            else:
                r['error'] = f'Lookup failed: {err}'

        return r

    def lookup_file(self, file_path: str) -> dict:
        """Compute hash from file then lookup."""
        hashes = HashVerifier.compute(file_path)
        if 'error' in hashes:
            return {'error': hashes['error']}
        return self.lookup(hashes['SHA256'])


# ═══════════════════════════════════════════════════════════════
#  16. Keyword Search in Files
# ═══════════════════════════════════════════════════════════════

class KeywordSearcher:
    """
    Recursively search a folder for keywords in file contents.
    Supports text files, binary files, PDFs (text layer only).
    Returns file path, line number, matched line.
    """

    TEXT_EXTENSIONS = {
        '.txt','.log','.csv','.xml','.json','.html','.htm',
        '.py','.js','.php','.sql','.bat','.ps1','.vbs','.sh',
        '.cfg','.ini','.conf','.yaml','.yml','.md','.rtf',
    }
    DOCX_EXTENSIONS = {'.docx', '.doc'}
    PDF_EXTENSIONS  = {'.pdf'}

    def __init__(self):
        self.results = []
        self.errors  = []
        self.stats   = {}

    def search(self, folder: str, keywords: list[str],
               case_sensitive: bool = False,
               max_results: int = 500) -> list:
        self.results = []
        self.errors  = []
        files_scanned = 0
        files_matched = 0

        if not os.path.isdir(folder):
            self.errors.append(f'Folder not found: {folder}')
            return []

        kws = keywords if case_sensitive else [k.lower() for k in keywords]

        for root, dirs, files in os.walk(folder):
            # Skip system dirs
            dirs[:] = [d for d in dirs if d not in
                       ('$Recycle.Bin', 'System Volume Information', '.git',
                        '__pycache__', 'node_modules')]
            for fn in files:
                if len(self.results) >= max_results:
                    break
                fp  = os.path.join(root, fn)
                ext = Path(fp).suffix.lower()
                files_scanned += 1

                try:
                    if ext in self.TEXT_EXTENSIONS:
                        matched = self._search_text(fp, kws, case_sensitive)
                    elif ext in self.DOCX_EXTENSIONS:
                        matched = self._search_docx(fp, kws, case_sensitive)
                    elif ext in self.PDF_EXTENSIONS:
                        matched = self._search_pdf(fp, kws, case_sensitive)
                    else:
                        matched = self._search_binary(fp, kws, case_sensitive)

                    if matched:
                        files_matched += 1
                        self.results.extend(matched)
                except Exception as e:
                    self.errors.append(f'{fp}: {e}')

        self.stats = {
            'files_scanned': files_scanned,
            'files_matched': files_matched,
            'total_hits'   : len(self.results),
            'folder'       : folder,
            'keywords'     : keywords,
        }
        return self.results

    def _search_text(self, fp: str, kws: list, case_sensitive: bool) -> list:
        hits = []
        with open(fp, 'r', encoding='utf-8', errors='replace') as f:
            for lineno, line in enumerate(f, 1):
                check = line if case_sensitive else line.lower()
                for kw in kws:
                    if kw in check:
                        hits.append({
                            'file'   : fp,
                            'line_no': lineno,
                            'keyword': kw,
                            'context': line.strip()[:120],
                            'type'   : 'text',
                        })
                        break
        return hits

    def _search_docx(self, fp: str, kws: list, case_sensitive: bool) -> list:
        hits = []
        try:
            import docx
            doc = docx.Document(fp)
            lines = []
            for para in doc.paragraphs:
                if para.text.strip():
                    lines.append(para.text)
            # Also check tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            lines.append(cell.text)
            for lineno, line in enumerate(lines, 1):
                check = line if case_sensitive else line.lower()
                for kw in kws:
                    if kw in check:
                        hits.append({
                            'file'   : fp,
                            'line_no': lineno,
                            'keyword': kw,
                            'context': line.strip()[:120],
                            'type'   : 'docx',
                        })
                        break
        except ImportError:
            self.errors.append('python-docx not installed: pip install python-docx')
        except Exception as e:
            self.errors.append(f'DOCX read error ({fp}): {e}')
        return hits

    def _search_pdf(self, fp: str, kws: list, case_sensitive: bool) -> list:
        hits = []
        try:
            import fitz   # PyMuPDF
            doc = fitz.open(fp)
            for page_num, page in enumerate(doc, 1):
                text = page.get_text()
                if not text.strip():
                    continue
                lines = text.split('\n')
                for line in lines:
                    check = line if case_sensitive else line.lower()
                    for kw in kws:
                        if kw in check:
                            hits.append({
                                'file'   : fp,
                                'line_no': f'Page {page_num}',
                                'keyword': kw,
                                'context': line.strip()[:120],
                                'type'   : 'pdf',
                            })
                            break
            doc.close()
        except ImportError:
            # Fallback: try pdfplumber or basic text extraction
            try:
                with open(fp, 'rb') as f:
                    raw = f.read()
                text = raw.decode('latin-1', errors='replace')
                check = text if case_sensitive else text.lower()
                for kw in kws:
                    if kw in check:
                        idx = check.find(kw)
                        ctx = text[max(0,idx-40):idx+80].replace('\n',' ')
                        hits.append({
                            'file': fp, 'line_no': 'N/A',
                            'keyword': kw, 'context': ctx[:120], 'type': 'pdf'
                        })
                        break
            except Exception:
                pass
        except Exception as e:
            self.errors.append(f'PDF read error ({fp}): {e}')
        return hits

    def _search_binary(self, fp: str, kws: list, case_sensitive: bool) -> list:
        hits = []
        try:
            with open(fp, 'rb') as f:
                data = f.read(5 * 1024 * 1024)  # first 5 MB only
            text = data.decode('utf-8', errors='replace')
            check = text if case_sensitive else text.lower()
            for kw in kws:
                if kw in check:
                    idx = check.find(kw)
                    context = text[max(0,idx-40):idx+80].replace('\n',' ').replace('\r','')
                    hits.append({
                        'file'   : fp,
                        'line_no': 'N/A',
                        'keyword': kw,
                        'context': context[:120],
                        'type'   : 'binary',
                    })
                    break
        except Exception:
            pass
        return hits


# ═══════════════════════════════════════════════════════════════
#  17. Timeline Chart Generator  (Matplotlib)
# ═══════════════════════════════════════════════════════════════

class TimelineChartGenerator:
    """
    Generates forensic timeline charts from Autopsy CSV data:
    - Bar chart: file activity per day/hour
    - Pie chart: file type distribution
    - Bar chart: suspicious vs clean files
    Saves as PNG files.
    """

    def __init__(self, records: list[dict], output_dir: str):
        self.records    = records
        self.output_dir = output_dir
        self.generated  = []
        os.makedirs(output_dir, exist_ok=True)

    def _try_parse_date(self, val: str):
        for fmt in ('%Y-%m-%d %H:%M:%S','%Y-%m-%dT%H:%M:%S',
                    '%m/%d/%Y %H:%M','%Y-%m-%d','%d/%m/%Y'):
            try:
                import datetime as dt
                return dt.datetime.strptime(str(val).strip(), fmt)
            except Exception:
                pass
        return None

    def generate_all(self) -> list:
        try:
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as plt
            import matplotlib.ticker as ticker
            from collections import Counter
        except ImportError:
            return ['matplotlib not installed.']

        # Dark forensic theme
        plt.rcParams.update({
            'figure.facecolor' : '#0d1117',
            'axes.facecolor'   : '#161b22',
            'axes.edgecolor'   : '#30363d',
            'axes.labelcolor'  : '#e6edf3',
            'xtick.color'      : '#8b949e',
            'ytick.color'      : '#8b949e',
            'text.color'       : '#e6edf3',
            'grid.color'       : '#30363d',
            'grid.linestyle'   : '--',
            'grid.alpha'       : 0.5,
        })

        # ── 1. Activity per day ───────────────────────────────
        date_candidates = ['Date','date','Timestamp','timestamp','Modified','Created']
        date_field = None
        for c in date_candidates:
            if self.records and c in self.records[0]:
                date_field = c; break

        if date_field:
            dates = []
            for rec in self.records:
                d = self._try_parse_date(rec.get(date_field,''))
                if d: dates.append(d.strftime('%Y-%m-%d'))

            if dates:
                counts = Counter(dates)
                sorted_dates = sorted(counts.keys())[-30:]  # last 30 days
                values = [counts[d] for d in sorted_dates]

                fig, ax = plt.subplots(figsize=(12, 5))
                bars = ax.bar(range(len(sorted_dates)), values,
                              color='#1f6feb', edgecolor='#58a6ff', linewidth=0.5)
                ax.set_xticks(range(len(sorted_dates)))
                ax.set_xticklabels(sorted_dates, rotation=45, ha='right', fontsize=7)
                ax.set_title('File Activity Timeline (Events per Day)',
                             fontsize=13, pad=12, color='#58a6ff', fontweight='bold')
                ax.set_ylabel('Number of Events')
                ax.set_xlabel('Date')
                ax.grid(axis='y')
                fig.tight_layout()
                path1 = os.path.join(self.output_dir, 'chart_activity_timeline.png')
                fig.savefig(path1, dpi=150, bbox_inches='tight', facecolor=fig.get_facecolor())
                plt.close(fig)
                self.generated.append(path1)

        # ── 2. File type distribution (pie) ──────────────────
        ext_field = None
        for c in ('Name','File Name','name','filename'):
            if self.records and c in self.records[0]:
                ext_field = c; break

        if ext_field:
            exts = []
            for rec in self.records:
                name = rec.get(ext_field,'')
                ext  = Path(str(name)).suffix.lower() or 'no ext'
                exts.append(ext)
            counts = Counter(exts).most_common(10)
            labels = [c[0] for c in counts]
            values = [c[1] for c in counts]

            colors = ['#1f6feb','#f85149','#3fb950','#d29922','#58a6ff',
                      '#bc8cff','#ff7b72','#79c0ff','#56d364','#ffa657']
            fig, ax = plt.subplots(figsize=(8, 6))
            wedges, texts, autotexts = ax.pie(
                values, labels=labels, autopct='%1.1f%%',
                colors=colors[:len(labels)], startangle=140,
                textprops={'color':'#e6edf3','fontsize':9})
            for at in autotexts: at.set_color('#0d1117'); at.set_fontweight('bold')
            ax.set_title('File Type Distribution', fontsize=13, pad=12,
                         color='#58a6ff', fontweight='bold')
            fig.tight_layout()
            path2 = os.path.join(self.output_dir, 'chart_file_types.png')
            fig.savefig(path2, dpi=150, bbox_inches='tight', facecolor='#0d1117')
            plt.close(fig)
            self.generated.append(path2)

        # ── 3. Suspicious vs Clean bar ────────────────────────
        total = len(self.records)
        sus_count = sum(
            1 for r in self.records
            if any(r.get(c,'').lower().endswith(ext)
                   for ext in ('.exe','.bat','.ps1','.vbs','.dll','.scr','.cmd')
                   for c in ('Name','File Name','name'))
        )
        clean = total - sus_count

        fig, ax = plt.subplots(figsize=(6, 5))
        bars = ax.bar(['Clean Files','Suspicious Files'], [clean, sus_count],
                      color=['#3fb950','#f85149'], edgecolor='#30363d', linewidth=0.8, width=0.4)
        for bar, val in zip(bars, [clean, sus_count]):
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5,
                    str(val), ha='center', va='bottom', color='#e6edf3', fontweight='bold')
        ax.set_title('Clean vs Suspicious Files', fontsize=13, pad=12,
                     color='#58a6ff', fontweight='bold')
        ax.set_ylabel('Number of Files')
        ax.grid(axis='y')
        fig.tight_layout()
        path3 = os.path.join(self.output_dir, 'chart_suspicious_vs_clean.png')
        fig.savefig(path3, dpi=150, bbox_inches='tight', facecolor=fig.get_facecolor())
        plt.close(fig)
        self.generated.append(path3)

        return self.generated


# ═══════════════════════════════════════════════════════════════
#  18. Excel Exporter
# ═══════════════════════════════════════════════════════════════

class ExcelExporter:
    """Export forensic analysis results to a styled .xlsx workbook."""

    def __init__(self, output_path: str):
        self.output_path = output_path
        self.workbook    = None

    def export(self, data: dict) -> bool:
        """
        data = {
          'Timeline'  : list of dicts,
          'Suspicious': list of dicts,
          'Metadata'  : dict,
          'Hashes'    : dict,
          'Browser'   : list of dicts,
          ...
        }
        """
        try:
            import openpyxl
            from openpyxl.styles import (PatternFill, Font, Alignment,
                                          Border, Side)

            wb = openpyxl.Workbook()
            wb.remove(wb.active)  # remove default sheet

            # Colour palette
            H_FILL  = PatternFill('solid', fgColor='1F6FEB')   # blue header
            ALT_FILL= PatternFill('solid', fgColor='1C2128')   # alt row
            SUS_FILL= PatternFill('solid', fgColor='3D1A1A')   # suspicious row
            H_FONT  = Font(name='Calibri', bold=True, color='FFFFFF', size=10)
            B_FONT  = Font(name='Calibri', size=9, color='E6EDF3')
            CENTER  = Alignment(horizontal='center', vertical='center', wrap_text=True)
            LEFT    = Alignment(horizontal='left',   vertical='center', wrap_text=True)

            thin = Side(style='thin', color='30363D')
            border = Border(left=thin, right=thin, top=thin, bottom=thin)

            def add_sheet(name, rows: list[dict], sus_keys: set = None):
                if not rows: return
                ws = wb.create_sheet(title=name[:31])
                cols = list(rows[0].keys())
                cols = [c for c in cols if not c.startswith('_')]

                # Header row
                for ci, col in enumerate(cols, 1):
                    cell = ws.cell(row=1, column=ci, value=col)
                    cell.fill = H_FILL; cell.font = H_FONT
                    cell.alignment = CENTER; cell.border = border
                    ws.column_dimensions[
                        openpyxl.utils.get_column_letter(ci)].width = 20

                # Data rows
                for ri, row in enumerate(rows, 2):
                    is_sus = sus_keys and any(
                        str(row.get(k,'')).lower().endswith(e)
                        for k in ('Name','File Name','name')
                        for e in sus_keys)
                    fill = SUS_FILL if is_sus else (ALT_FILL if ri%2==0 else None)
                    for ci, col in enumerate(cols, 1):
                        val = str(row.get(col,''))[:200]
                        cell = ws.cell(row=ri, column=ci, value=val)
                        cell.font   = Font(name='Calibri', size=9,
                                           color='F85149' if is_sus else 'E6EDF3')
                        cell.alignment = LEFT
                        cell.border    = border
                        if fill: cell.fill = fill

                ws.freeze_panes = 'A2'

            def add_kv_sheet(name, kv: dict):
                ws = wb.create_sheet(title=name[:31])
                ws.column_dimensions['A'].width = 25
                ws.column_dimensions['B'].width = 60
                ws.cell(1,1,'Property').fill=H_FILL; ws.cell(1,1).font=H_FONT
                ws.cell(1,2,'Value').fill=H_FILL;    ws.cell(1,2).font=H_FONT
                for ri,(k,v) in enumerate(kv.items(),2):
                    ws.cell(ri,1,str(k)).font = Font(bold=True,size=9,color='58A6FF')
                    ws.cell(ri,2,str(v)).font  = Font(size=9,color='E6EDF3')
                    if ri%2==0:
                        ws.cell(ri,1).fill=ALT_FILL
                        ws.cell(ri,2).fill=ALT_FILL

            # Sheet: Timeline
            if 'Timeline' in data and data['Timeline']:
                sus_exts = {'.exe','.bat','.ps1','.vbs','.dll','.scr','.cmd'}
                add_sheet('Timeline', data['Timeline'], sus_exts)

            # Sheet: Suspicious
            if 'Suspicious' in data and data['Suspicious']:
                add_sheet('Suspicious Files', data['Suspicious'])

            # Sheet: Metadata
            if 'Metadata' in data and data['Metadata']:
                flat = {k:v for k,v in data['Metadata'].items()
                        if k != 'exif' and not isinstance(v,dict)}
                exif = data['Metadata'].get('exif',{})
                flat.update({f'EXIF_{k}':v for k,v in exif.items()
                             if not isinstance(v,dict)})
                add_kv_sheet('File Metadata', flat)

            # Sheet: Browser
            if 'Browser' in data and data['Browser']:
                add_sheet('Browser History', data['Browser'])

            # Sheet: Summary
            ws_sum = wb.create_sheet(title='Summary', index=0)
            ws_sum.column_dimensions['A'].width = 30
            ws_sum.column_dimensions['B'].width = 40
            summary_rows = [
                ('ForensiTrace Report',''),
                ('Generated At', datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
                ('Submitted To','Engr. Rafia Durrani'),
                ('Submitted By','Ushna Hidayat | Noor Fatima'),
                ('Course','CSDF-30117 Introduction to Digital Forensics'),
                ('',''),
            ]
            for k,v in data.items():
                if isinstance(v, list): summary_rows.append((f'{k} Records', str(len(v))))
            for ri,(k,v) in enumerate(summary_rows,1):
                ws_sum.cell(ri,1,k).font = Font(bold=True,color='58A6FF',size=10)
                ws_sum.cell(ri,2,v).font = Font(size=10,color='E6EDF3')

            # Set workbook background
            for ws in wb.worksheets:
                ws.sheet_properties.tabColor = '1F6FEB'

            wb.save(self.output_path)
            return True
        except ImportError:
            return False
        except Exception as e:
            print(f'Excel export error: {e}')
            return False